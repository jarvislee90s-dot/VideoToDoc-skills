from __future__ import annotations

import os
import re
from pathlib import Path

from .config import Settings
from .io import write_text
from .models import Section


def render_markdown(title: str, sections: list[Section], mindmap: str, output_path: Path) -> Path:
    """兼容旧调用：生成保留原始转录换行的 Markdown。"""

    return render_original_markdown(title, sections, output_path)


def render_original_markdown(title: str, sections: list[Section], output_path: Path) -> Path:
    lines = [f"# {title}", "", "## 图文讲义", ""]
    for section in sections:
        lines.extend(
            [
                f"### 第 {section.slide_index} 页",
                "",
                f"时间：{_format_ms(section.start_ms)} - {_format_ms(section.end_ms)}",
                "",
                f"![第 {section.slide_index} 页]({_markdown_image_path(section.image_path, output_path.parent)})",
                "",
                section.transcript.strip() or "本页无讲解。",
                "",
            ]
        )
        if section.notes:
            lines.extend([f"> 标记：{', '.join(section.notes)}", ""])
    write_text(output_path, "\n".join(lines).rstrip() + "\n")
    return output_path


def render_compact_markdown(
    title: str,
    sections: list[Section],
    output_path: Path,
    mindmap_image_path: Path | None = None,
) -> Path:
    lines = [f"# {title}", "", "## 图文讲义", ""]
    for index, section in enumerate(sections):
        lines.extend(
            [
                f"### 第 {section.slide_index} 页 · {_format_ms(section.start_ms)} - {_format_ms(section.end_ms)}",
                "",
                f"![第 {section.slide_index} 页]({_markdown_image_path(section.image_path, output_path.parent)})",
                "",
                _compact_transcript(section.transcript),
                "",
            ]
        )
        if section.notes:
            lines.extend([f"> 标记：{', '.join(section.notes)}", ""])
        if index != len(sections) - 1:
            lines.extend(["---", ""])
    if mindmap_image_path:
        lines.extend(["## 思维导图", "", f"![思维导图]({_markdown_image_path(mindmap_image_path, output_path.parent)})", ""])
    write_text(output_path, "\n".join(lines).rstrip() + "\n")
    return output_path


def ensure_semantic_markdown(
    title: str,
    sections: list[Section],
    output_path: Path,
    mindmap_image_path: Path | None = None,
) -> Path:
    if output_path.exists():
        return ensure_mindmap_link(output_path, mindmap_image_path)
    lines = [
        f"# {title}",
        "",
    ]
    for index, section in enumerate(sections):
        lines.extend(
            [
                f"### 第 {section.slide_index} 页 · {_format_ms(section.start_ms)} - {_format_ms(section.end_ms)}",
                "",
                f"<!-- IMAGE:{section.slide_index} -->",
                "",
                _compact_transcript(section.transcript),
                "",
            ]
        )
        if index != len(sections) - 1:
            lines.extend(["---", ""])
    write_text(output_path, "\n".join(lines).rstrip() + "\n")
    return ensure_mindmap_link(output_path, mindmap_image_path)


def ensure_mindmap_link(markdown_path: Path, mindmap_image_path: Path | None) -> Path:
    if not mindmap_image_path:
        return markdown_path
    image_ref = _markdown_image_path(mindmap_image_path, markdown_path.parent)
    text = markdown_path.read_text(encoding="utf-8") if markdown_path.exists() else ""
    replacement = f"![思维导图]({image_ref})"
    if re.search(r"!\[思维导图\]\([^)]+\)", text):
        text = re.sub(r"!\[思维导图\]\([^)]+\)", replacement, text)
    elif "## 思维导图" in text:
        text = text.rstrip() + f"\n\n{replacement}\n"
    else:
        text = text.rstrip() + f"\n\n## 思维导图\n\n{replacement}\n"
    write_text(markdown_path, text.rstrip() + "\n")
    return markdown_path


def markdown_to_docx(markdown_path: Path, output_path: Path) -> Path | None:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
    except ModuleNotFoundError:
        return None

    document = Document()
    _configure_document_styles(document)
    base_dir = markdown_path.parent
    pending: list[str] = []

    def flush_pending() -> None:
        if not pending:
            return
        text = " ".join(item.strip() for item in pending if item.strip())
        pending.clear()
        if text:
            document.add_paragraph(text, style="LectureBody")

    for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            flush_pending()
            continue
        if line == "---":
            flush_pending()
            _add_separator(document)
            continue
        if line.startswith("# "):
            flush_pending()
            para = document.add_paragraph(style="LectureTitle")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(line[2:].strip())
            continue
        if line.startswith("## "):
            flush_pending()
            document.add_paragraph(line[3:].strip(), style="LectureHeading")
            continue
        if line.startswith("### "):
            flush_pending()
            para = document.add_paragraph(style="LecturePageHeading")
            run = para.add_run(line[4:].strip())
            if "·" in run.text:
                page, _, time_range = run.text.partition("·")
                run.text = page.strip()
                time_run = para.add_run(f"  {time_range.strip()}")
                time_run.font.color.rgb = RGBColor(90, 102, 122)
                time_run.font.size = Pt(10.5)
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            flush_pending()
            image_path = _resolve_markdown_path(image_match.group(2), base_dir)
            if image_path.exists():
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                width = Inches(6.4 if image_match.group(1) != "思维导图" else 6.8)
                run.add_picture(str(image_path), width=width)
            continue
        if line.startswith("> "):
            flush_pending()
            document.add_paragraph(line[2:].strip(), style="LectureNote")
            continue
        if line.startswith("- "):
            flush_pending()
            # List items get their own paragraph with bullet styling
            para = document.add_paragraph(style="List Bullet")
            # Remove the "- " prefix and handle bold text
            text = line[2:].strip()
            # Handle **bold** text
            runs = re.split(r'(\*\*[^*]+\*\*)', text)
            for run in runs:
                if run.startswith('**') and run.endswith('**'):
                    run_obj = para.add_run(run[2:-2])
                    run_obj.bold = True
                elif run:
                    para.add_run(run)
            continue
        pending.append(line)
    flush_pending()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def generate_mindmap(title: str, sections: list[Section], output_path: Path, settings: Settings | None = None) -> str:
    settings = settings or Settings()
    if settings.mindmap_backend in {"auto", "llm"} and os.getenv("OPENAI_API_KEY"):
        try:
            mindmap = _generate_mindmap_with_openai(title, sections, settings)
            write_text(output_path, mindmap)
            return mindmap
        except Exception:
            if settings.mindmap_backend == "llm":
                raise

    lines = ["mindmap", f"  root(({_safe_mermaid(title)}))"]
    for section in sections[:20]:
        heading = _extract_heading(section.transcript, section.slide_index)
        lines.append(f"    第 {section.slide_index} 页")
        lines.append(f"      {_safe_mermaid(heading)}")
    if len(sections) > 20:
        lines.append("    更多页面")
        lines.append(f"      共 {len(sections)} 页，详见正文")
    mindmap = "\n".join(lines)
    write_text(output_path, mindmap)
    return mindmap


def render_docx(title: str, sections: list[Section], mindmap: str, output_path: Path) -> Path | None:
    compact_path = output_path.with_name("draft_compact.md")
    render_compact_markdown(title, sections, compact_path)
    return markdown_to_docx(compact_path, output_path)


def _configure_document_styles(document: object) -> None:
    from docx.enum.style import WD_STYLE_TYPE  # type: ignore
    from docx.enum.text import WD_LINE_SPACING  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    def style(name: str, base: str = "Normal"):
        if name in styles:
            return styles[name]
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    title = style("LectureTitle")
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(12)

    heading = style("LectureHeading")
    heading.font.name = "Microsoft YaHei"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)

    page_heading = style("LecturePageHeading")
    page_heading.font.name = "Microsoft YaHei"
    page_heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    page_heading.font.size = Pt(12.5)
    page_heading.font.bold = True
    page_heading.font.color.rgb = RGBColor(28, 96, 175)
    page_heading.paragraph_format.space_before = Pt(8)
    page_heading.paragraph_format.space_after = Pt(6)
    page_heading.paragraph_format.keep_with_next = True

    body = style("LectureBody")
    body.font.name = "Microsoft YaHei"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    body.font.size = Pt(10.5)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(5.25)

    note = style("LectureNote")
    note.font.name = "Microsoft YaHei"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor(90, 102, 122)
    note.paragraph_format.space_after = Pt(5.25)


def _add_separator(document: object) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Pt  # type: ignore

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    p_pr.append(border)


def _compact_transcript(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "本页无讲解。"
    paragraphs: list[str] = []
    buffer: list[str] = []
    length = 0
    for line in lines:
        buffer.append(line)
        length += len(line)
        if length >= 90 and re.search(r"[。！？?!]$", line):
            paragraphs.append("".join(buffer))
            buffer = []
            length = 0
        elif length >= 180:
            paragraphs.append("".join(buffer))
            buffer = []
            length = 0
    if buffer:
        paragraphs.append("".join(buffer))
    return "\n\n".join(paragraphs)


def _markdown_image_path(image_path: str | Path, output_dir: Path) -> str:
    path = Path(image_path)
    if not path.is_absolute():
        absolute = (Path.cwd() / path).resolve()
    else:
        absolute = path.resolve()
    try:
        return absolute.relative_to(output_dir.resolve()).as_posix()
    except ValueError:
        return str(image_path)


def _resolve_markdown_path(value: str, base_dir: Path) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    candidate = base_dir / path
    if candidate.exists():
        return candidate
    return Path.cwd() / path


def _extract_heading(text: str, slide_index: int) -> str:
    cleaned = " ".join(text.replace("\n", " ").split())
    if not cleaned or cleaned == "本页无讲解。":
        return "本页无讲解"
    return cleaned[:36]


def _safe_mermaid(text: str) -> str:
    return text.replace("(", "（").replace(")", "）").replace(":", "：").strip() or "未命名"


def _format_ms(value: int) -> str:
    seconds = value // 1000
    minutes, sec = divmod(seconds, 60)
    hours, minute = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minute:02d}:{sec:02d}"
    return f"{minute:02d}:{sec:02d}"


def _generate_mindmap_with_openai(title: str, sections: list[Section], settings: Settings) -> str:
    from openai import OpenAI  # type: ignore

    body = "\n\n".join(
        f"第 {section.slide_index} 页：{section.transcript[:600]}" for section in sections[:80]
    )
    prompt = f"""请根据课程讲义内容生成 Mermaid mindmap。
要求：
1. 只输出 Mermaid mindmap 正文，不要 Markdown 代码围栏。
2. 根节点是课程标题。
3. 层级清晰，优先保留章节、概念、公式、风险点和例子。
4. 每个节点尽量短，方便在飞书文档里继续编辑。

课程标题：{title}

内容：
{body}
"""
    client = OpenAI()
    response = client.chat.completions.create(
        model=settings.mindmap_model,
        messages=[
            {"role": "system", "content": "你是课程讲义整理助手，擅长输出可编辑 Mermaid mindmap。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    content = response.choices[0].message.content or ""
    return _strip_code_fence(content.strip())


def _strip_code_fence(text: str) -> str:
    if text.startswith("```"):
        lines = text.splitlines()
        if lines:
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return text
